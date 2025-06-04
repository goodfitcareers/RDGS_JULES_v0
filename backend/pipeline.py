import json
import mimetypes
import os
from datetime import (
    date,
    datetime,
)
from typing import Any, Dict, List, Optional, TypedDict

# from .models import Client # If client_id needs to be validated against Client table as UUID
import uuid # Import uuid module
from uuid import UUID  # If client_id is expected to be UUID

from docx import Document
from langgraph.graph import END, StateGraph

# Ensure timezone is imported if used by parse_date_string
from openai import (
    APIConnectionError,
    APIStatusError,
    APITimeoutError,
    OpenAI,
    RateLimitError,
)

# Fallback imports
from pypdf import PdfReader
from sqlalchemy.exc import SQLAlchemyError
from sqlmodel import Session, create_engine

# Primary extractor
from unstructured.partition import auto # Changed import

from .models import (  # Assuming Role and RoleStatus are in models.py
    Role,
    RoleStatus,
)

# This will cause an ImportError if Role/RoleStatus are not in models.py
from .settings import settings  # For API keys, LangSmith settings, etc.

# Setup LangSmith tracing
if settings.LANGCHAIN_TRACING_V2:
    os.environ["LANGCHAIN_TRACING_V2"] = "true"
# LANGSMITH_API_KEY should be picked up automatically by LangChain if set in the environment.
# It's typically not accessed via settings for security reasons (should be in .env directly)
if settings.LANGSMITH_PROJECT:
    os.environ["LANGCHAIN_PROJECT"] = settings.LANGSMITH_PROJECT


class GraphState(TypedDict):
    client_id: Optional[str]
    source_document_paths: List[str]  # Paths to general source documents
    resume_file_path: Optional[str]  # Path to the final resume

    # Output fields from extract_node
    source_document_texts: List[str]
    resume_text: Optional[str]

    # Output from parse_roles_node
    roles_draft: Optional[List[Dict[str, Any]]]

    # Output from sync_roles_node
    synced_roles_count: Optional[int]

    # Error/status tracking
    error_message: Optional[str]  # Accumulates errors from all nodes
    current_node_error: Optional[str]  # For errors specific to the last run node
    trace_id: Optional[str] # For LangSmith tracing


def extract_text_from_file(file_path: str, current_trace_id: Optional[str] = None) -> str: # Added trace_id for logging
    # Note: The print statements within this utility function are NOT modified to include trace_id
    # as it's a generic utility. Trace ID logging is primarily for the main graph nodes.
    # If detailed tracing within this function was needed, current_trace_id would be passed and used.
    if not os.path.exists(file_path):
        print(f"Error: File not found at {file_path}") # Not prepending here
        raise FileNotFoundError(f"File not found: {file_path}")

    mime_type, _ = mimetypes.guess_type(file_path)
    text_content = ""
    extraction_method_used = "None"

    try:
        # Attempt 1: Use unstructured.io
        # Note: unstructured[local-inference] was removed due to space issues in previous setup.
        # This means partition_file might only work for basic types like .txt, .md, .html
        # unless it has other built-in parsers that don't rely on detectron2/torch.
        # Not prepending trace_id to these print statements
        print(
            f"Attempting extraction with unstructured.io for {file_path} (MIME: {mime_type})"
        )
        elements = auto.partition_file(filename=file_path) # Changed call
        text_content = "\n\n".join([el.text for el in elements if hasattr(el, "text")])
        if text_content.strip():
            extraction_method_used = "unstructured.io"
            print(f"Successfully extracted text from {file_path} using unstructured.io")
            return text_content
        else:
            print(
                f"unstructured.io extracted no content from {file_path}. Proceeding to fallbacks."
            )
    except Exception as e:
        print(
            f"unstructured.io failed for {file_path} (MIME: {mime_type}): {e}. Attempting fallback."
        )

    # Attempt 2: Fallbacks
    try:
        print(f"Attempting fallback extraction for {file_path} (MIME: {mime_type})") # Not prepending
        if mime_type == "application/pdf":
            with open(file_path, "rb") as f:
                reader = PdfReader(f)
                # Concatenate text from all pages, ensuring None returns from extract_text are handled
                page_texts = [page.extract_text() for page in reader.pages]
                text_content = "".join(
                    pt for pt in page_texts if pt
                )  # Filter out None before join
            extraction_method_used = "pypdf"
        elif (
            mime_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            doc = Document(file_path)
            text_content = "\n".join([para.text for para in doc.paragraphs])
            extraction_method_used = "python-docx"
        elif mime_type in ["text/markdown", "text/plain", "text/html"] or (
            mime_type is None and file_path.lower().endswith((".md", ".txt", ".html"))
        ):
            # Adding .html here as unstructured might not handle it well without all extras.
            # Also handling cases where MIME might be None but extension is clear.
            with open(file_path, encoding="utf-8", errors="ignore") as f:
                text_content = f.read()
            extraction_method_used = "direct_read"
        else:
            print( # Not prepending
                f"Unsupported file type for fallback: {mime_type} for file {file_path}"
            )
            # Not raising an error here, but returning empty string, node will report aggregate errors
            return ""

        if text_content.strip():
            print( # Not prepending
                f"Successfully extracted text from {file_path} using fallback: {extraction_method_used}."
            )
        else:
            print( # Not prepending
                f"Fallback ({extraction_method_used}) extracted no content from {file_path}."
            )
        return text_content
    except Exception as e:
        print(f"Fallback extraction failed for {file_path}: {e}") # Not prepending
        # Raise a specific error that can be caught by the node
        raise OSError(
            f"Failed to extract text from {file_path} using all methods (MIME: {mime_type}). Last error: {e}"
        ) from e


def extract_node(state: GraphState) -> Dict[str, Any]:
    current_trace_id = state.get("trace_id") or str(uuid.uuid4())
    print(f"[Trace ID: {current_trace_id}] --- Running Extract Node ---")
    source_texts: List[str] = []
    resume_text: Optional[str] = (
        None  # Initialize to ensure it's always in the output dict
    )
    current_node_error_messages: List[str] = []

    try:
        resume_file_path = state.get("resume_file_path")
        if resume_file_path:  # Check if the key exists and is not None/empty
            print(f"[Trace ID: {current_trace_id}] Extracting text from resume: {resume_file_path}")
            try:
                resume_text = extract_text_from_file(resume_file_path, current_trace_id)
            except Exception as e:
                msg = f"Failed to extract resume '{resume_file_path}': {str(e)}"
                print(f"[Trace ID: {current_trace_id}] {msg}")
                current_node_error_messages.append(msg)
        else:
            print(f"[Trace ID: {current_trace_id}] No resume_file_path provided or it's empty in state.")
            # Depending on requirements, this might be an error or acceptable
            # current_node_error_messages.append("Resume file path was not provided.")

        source_doc_paths = state.get("source_document_paths", [])
        print(f"[Trace ID: {current_trace_id}] Extracting text from {len(source_doc_paths)} source document(s).")
        for doc_path in source_doc_paths:
            print(f"[Trace ID: {current_trace_id}] Extracting from source document: {doc_path}")
            try:
                text = extract_text_from_file(doc_path, current_trace_id)
                source_texts.append(text)
            except Exception as e:
                msg = f"Failed to extract source document '{doc_path}': {str(e)}"
                print(f"[Trace ID: {current_trace_id}] {msg}")
                current_node_error_messages.append(msg)
                source_texts.append(
                    ""
                )  # Add empty string for failed extractions to maintain list length

    except Exception as e:
        # Catch-all for unexpected errors within the node logic itself
        print(f"[Trace ID: {current_trace_id}] Unexpected critical error in extract_node: {str(e)}")
        current_node_error_messages.append(
            f"A critical unexpected error occurred during text extraction: {str(e)}"
        )

    # Consolidate error messages for this node's execution
    node_error_output = (
        "; ".join(current_node_error_messages) if current_node_error_messages else None
    )

    # Error accumulation logic:
    # The 'error_message' key in the state should accumulate errors from all nodes.
    # 'current_node_error' is specific to the current node's execution.
    # If a node runs successfully, it should clear its 'current_node_error'.

    # Start with existing accumulated errors
    final_accumulated_error = state.get("error_message")

    if node_error_output:  # If this node had errors
        if final_accumulated_error:  # And there were previous errors
            final_accumulated_error = f"{final_accumulated_error}; {node_error_output}"
        else:  # No previous errors, just this node's errors
            final_accumulated_error = node_error_output

    # Return full state, plus outputs of this node
    # It's important that GraphState TypedDict keys are all potentially present
    # even if their values are None or empty lists.
    # LangGraph merges this dict back into the main state.
    return {
        "trace_id": current_trace_id,
        "source_document_texts": source_texts,
        "resume_text": resume_text,
        "error_message": final_accumulated_error,
        "current_node_error": node_error_output,
    }


def parse_roles_node(state: GraphState) -> Dict[str, Any]:
    current_trace_id = state.get("trace_id") or str(uuid.uuid4())
    print(f"[Trace ID: {current_trace_id}] --- Running Parse Roles Node ---")

    current_node_error_output: Optional[str] = None
    roles_draft_output: List[Dict[str, Any]] = []  # Initialize to empty list

    # If there was a critical error in a *previous* node that should stop processing,
    # or if resume_text is missing.
    # current_node_error is for this node's potential errors from previous runs if graph retries.
    # error_message is the accumulated error from all previous nodes.
    if state.get("error_message") and not state.get("current_node_error"):
        print(
            f"[Trace ID: {current_trace_id}] Skipping Parse Roles Node due to pre-existing critical errors from other nodes."
        )
        return {
            "trace_id": current_trace_id,
            "roles_draft": roles_draft_output,
            "current_node_error": None,
        }  # Return empty list, no new error from this node

    resume_text = state.get("resume_text")
    if not resume_text:
        print(f"[Trace ID: {current_trace_id}] No resume text to parse for roles.")
        current_node_error_output = "Cannot parse roles, resume text is missing."
        # roles_draft_output is already an empty list
    else:
        print(
            f"[Trace ID: {current_trace_id}] Parsing roles from resume text (length: {len(resume_text)}) using model: {settings.OPENAI_MODEL_NAME}"
        )

        client = OpenAI(api_key=settings.OPENAI_API_KEY)

        prompt = f"""
Extract the professional roles from the following resume text.
Return the output as a JSON object containing a single key "roles", which is a list of role objects.
Each role object should have the following keys:
- "company_name": string (name of the company)
- "title": string (job title)
- "start_date": string (e.g., "YYYY-MM-DD", "Mon YYYY", or "YYYY")
- "end_date": string (e.g., "YYYY-MM-DD", "Mon YYYY", "YYYY", or "Present")
- "description_points": list of strings (bullet points or paragraphs describing responsibilities and achievements)

Resume Text:
---
{resume_text}
---

JSON Output:
"""
        response_content_for_error_logging = (
            ""  # To store response content for logging in case of JSONDecodeError
        )
        try:
            response = client.chat.completions.create(
                model=settings.OPENAI_MODEL_NAME,
                messages=[{"role": "user", "content": prompt}],
                response_format={"type": "json_object"},
                temperature=0, # Added for deterministic output
                # Potential other parameters: max_tokens, etc.
            )

            response_content = response.choices[0].message.content
            response_content_for_error_logging = (
                response_content  # Store for potential error logging
            )
            if response_content:
                parsed_json = json.loads(response_content)
                if "roles" in parsed_json and isinstance(parsed_json["roles"], list):
                    roles_draft_output = parsed_json["roles"]
                    print(
                        f"[Trace ID: {current_trace_id}] Successfully parsed {len(roles_draft_output)} roles from LLM."
                    )
                else:
                    current_node_error_output = "LLM output is valid JSON but does not contain a 'roles' list or 'roles' is not a list."
                    print(f"[Trace ID: {current_trace_id}] {current_node_error_output}")
            else:
                current_node_error_output = "LLM response content is empty."
                print(f"[Trace ID: {current_trace_id}] {current_node_error_output}")

        except APITimeoutError as e:
            current_node_error_output = f"OpenAI API timeout: {str(e)}"
        except APIConnectionError as e:
            current_node_error_output = f"OpenAI API connection error: {str(e)}"
        except APIStatusError as e:
            current_node_error_output = f"OpenAI API status error (code {e.status_code}): {e.response.text if hasattr(e, 'response') and e.response else str(e)}"
        except RateLimitError as e:
            current_node_error_output = f"OpenAI API rate limit error: {str(e)}"
        except OpenAI.BadRequestError as e:  # Corrected namespace
            current_node_error_output = f"OpenAI API Bad Request (likely malformed JSON in LLM response or prompt issue): {str(e)}"
        except json.JSONDecodeError as e:
            current_node_error_output = f"Failed to decode JSON from LLM response: {str(e)}. Response content snippet: {response_content_for_error_logging[:500]}..."
        except Exception as e:  # Catch-all for other unexpected errors
            current_node_error_output = (
                f"An unexpected error occurred during LLM call or processing: {str(e)}"
            )

        if current_node_error_output:
            print(f"[Trace ID: {current_trace_id}] {current_node_error_output}")  # Print error specific to this attempt
            roles_draft_output = []  # Ensure output is empty list on error

    # Accumulate errors for the graph state
    final_accumulated_error = state.get("error_message")
    if current_node_error_output:
        if final_accumulated_error:  # If there were errors from previous nodes
            final_accumulated_error = (
                f"{final_accumulated_error}; {current_node_error_output}"
            )
        else:  # This node is the first to report an error
            final_accumulated_error = current_node_error_output
    # If this node ran successfully but there were prior errors, preserve them.
    # If this node ran successfully and there were no prior errors, final_accumulated_error remains None.

    return {
        "trace_id": current_trace_id,
        "roles_draft": roles_draft_output,
        "error_message": final_accumulated_error,
        "current_node_error": current_node_error_output,  # Specific error from this run
    }


# Date parsing helper function
def parse_date_string(date_str: Optional[str]) -> Optional[date]:
    if not date_str or date_str.strip().lower() == "present":
        return None

    # Try parsing "YYYY-MM-DD"
    try:
        return datetime.strptime(date_str.strip(), "%Y-%m-%d").date()
    except ValueError:
        pass

    # Try parsing "YYYY-MM" (default to day 1)
    try:
        return datetime.strptime(date_str.strip(), "%Y-%m").date().replace(day=1)
    except ValueError:
        pass

    # Try parsing "Month YYYY" e.g. "Jan 2020", "January 2020"
    try:
        return datetime.strptime(date_str.strip(), "%b %Y").date().replace(day=1)
    except ValueError:
        pass
    try:
        return datetime.strptime(date_str.strip(), "%B %Y").date().replace(day=1)
    except ValueError:
        pass

    # Try parsing "YYYY" (default to Jan 1st)
    try:
        return datetime.strptime(date_str.strip(), "%Y").date().replace(month=1, day=1)
    except ValueError:
        pass

    # Not prepending trace_id here as it's a utility function, and date parsing errors are more of data validation.
    print(
        f"Warning: Could not parse date string: '{date_str}' into a known date format (YYYY-MM-DD, YYYY-MM, Mon YYYY, YYYY). Returning None."
    )
    return None


def sync_roles_node(state: GraphState) -> Dict[str, Any]:
    current_trace_id = state.get("trace_id") or str(uuid.uuid4())
    print(f"[Trace ID: {current_trace_id}] --- Running Sync Roles Node ---")

    # Initialize outputs for this node
    current_node_error_output: Optional[str] = None
    synced_roles_count = 0

    # Pre-checks
    if state.get("error_message") and not state.get(
        "current_node_error"
    ):  # Check for errors from *previous* nodes
        print(
            f"[Trace ID: {current_trace_id}] Skipping Sync Roles Node due to pre-existing critical errors in the graph."
        )
        # Preserve existing error_message, set current_node_error to None, and return synced_roles_count 0
        return {
            "trace_id": current_trace_id,
            "synced_roles_count": 0,
            "error_message": state.get("error_message"),
            "current_node_error": None,
        }

    roles_to_sync = state.get("roles_draft")
    # Initialize final_accumulated_error with any pre-existing error message from the state.
    # This ensures that if the node exits early (e.g. no roles to sync), the original error isn't lost.
    final_accumulated_error = state.get("error_message")

    if not roles_to_sync:  # Handles None or empty list
        msg = "No roles draft provided or roles draft is empty; nothing to sync."
        print(f"[Trace ID: {current_trace_id}] {msg}")
        # Not an error for this node if previous node correctly produced no roles.
        # current_node_error_output remains None.
        # final_accumulated_error (from previous nodes, if any) is preserved.
        return {
            "trace_id": current_trace_id,
            "synced_roles_count": 0,
            "error_message": final_accumulated_error,
            "current_node_error": None,
        }

    client_id_str = state.get("client_id")
    if not client_id_str:
        current_node_error_output = "Client ID is missing in state, cannot sync roles."
        print(f"[Trace ID: {current_trace_id}] {current_node_error_output}")
    else:
        # Assuming client_id in the state is a string that needs to be UUID
        # This depends on how Client.id is defined in models.py
        # Role.client_id is a UUID.
        client_id_for_role: Optional[UUID] = None
        try:
            client_id_for_role = UUID(client_id_str)
        except ValueError:
            current_node_error_output = (
                f"Invalid client_id format: '{client_id_str}'. Must be a valid UUID."
            )
            print(f"[Trace ID: {current_trace_id}] {current_node_error_output}")
            # client_id_for_role remains None, sync will be skipped

        if (
            client_id_for_role and not current_node_error_output
        ):  # Proceed if client_id is valid UUID
            print(
                f"[Trace ID: {current_trace_id}] Attempting to sync {len(roles_to_sync)} roles for client_id: {client_id_for_role}"
            )

            engine = create_engine(settings.DATABASE_URL)  # Create engine here

            try:
                with Session(engine) as session:
                    for role_data in roles_to_sync:
                        # Basic check for essential fields from LLM output
                        if not role_data.get("company_name") or not role_data.get(
                            "title"
                        ):
                            print(
                                f"[Trace ID: {current_trace_id}] Skipping role due to missing company or title: {role_data}"
                            )
                            continue

                        # Map data to Role model instance
                        new_role = Role(
                            client_id=client_id_for_role,  # Uncommented and assuming client_id_for_role is UUID
                            company_name=role_data.get("company_name"),
                            title=role_data.get("title"),
                            start_date=parse_date_string(role_data.get("start_date")),
                            end_date=parse_date_string(role_data.get("end_date")),
                            # Assuming description_points from LLM becomes output_text
                            output_text="\n".join(
                                role_data.get("description_points", [])
                            ),
                            status=RoleStatus.PARSED,  # From your spec
                            # `id` will be auto-generated by DB
                            # `created_at`, `updated_at`, `revision` should have defaults in Role model
                            # `input_text_compact`, `validation_notes` are Optional, default to None
                        )
                        session.add(new_role)

                    session.commit()
                    synced_roles_count = len(roles_to_sync)  # Or count successful adds
                    print(
                        f"[Trace ID: {current_trace_id}] Successfully synced {synced_roles_count} roles to DB for client_id: {client_id_for_role}."
                    )
                    current_node_error_output = None  # Clear errors on success

            except SQLAlchemyError as e:
                msg = f"Database error while syncing roles for client {client_id_for_role}: {str(e)}"
                print(f"[Trace ID: {current_trace_id}] {msg}")
                current_node_error_output = msg
                # session.rollback() is implicitly handled by 'with Session' context manager on exception
            except (
                Exception
            ) as e:  # Catch any other unexpected errors during DB interaction
                msg = f"Unexpected error during database sync for client {client_id_for_role}: {str(e)}"
                print(f"[Trace ID: {current_trace_id}] {msg}")
                current_node_error_output = msg

    # Accumulate current_node_error_output into final_accumulated_error
    if current_node_error_output:
        if final_accumulated_error:
            final_accumulated_error = (
                f"{final_accumulated_error}; {current_node_error_output}"
            )
        else:
            final_accumulated_error = current_node_error_output

    return {
        "trace_id": current_trace_id,
        "synced_roles_count": synced_roles_count, # Always include synced_roles_count
        "error_message": final_accumulated_error,
        "current_node_error": current_node_error_output,
    }


# Graph Builder
workflow = StateGraph(GraphState)
workflow.add_node("extract", extract_node)
workflow.add_node("parse_roles", parse_roles_node)
workflow.add_node("sync_roles", sync_roles_node)

workflow.set_entry_point("extract")

# Define conditional edges based on 'current_node_error'
# This allows the graph to proceed to the next step only if the current one was successful.
# A more sophisticated error handling strategy might involve a dedicated error_handler_node.


def decide_next_node(state: GraphState) -> str:
    # Fetches trace_id for logging, but decision logic doesn't change based on it.
    # If trace_id is not set yet (e.g. error in the very first node before trace_id is set in its return),
    # this part might not have it, but the node itself would have logged with a generated one.
    current_trace_id = state.get("trace_id", "N/A_in_decide_next_node") # Default if not found
    if state.get("current_node_error"):
        print(
            f"[Trace ID: {current_trace_id}] Error detected: {state['current_node_error']}. Ending graph or routing to error handler."
        )
        return END  # Or an error handling node name
    # Check specific conditions for routing if needed
    print(f"[Trace ID: {current_trace_id}] No error in current node, proceeding.")
    return "continue"  # A conventional name for the "success" path


# Edges from extract_node
workflow.add_conditional_edges(
    "extract",
    decide_next_node,
    {
        "continue": "parse_roles",
        END: END,
    },  # If error, go to END. If success, go to parse_roles.
)

# Edges from parse_roles_node
workflow.add_conditional_edges(
    "parse_roles", decide_next_node, {"continue": "sync_roles", END: END}
)

# Edge from sync_roles_node (always goes to END or could have its own decision)
workflow.add_conditional_edges(
    "sync_roles",
    decide_next_node,  # Or simply END if no further conditional logic from sync
    {"continue": END, END: END},
)


app = workflow.compile()

# Example usage (for testing locally if needed)
# if __name__ == "__main__":
#     # Create dummy files for testing
#     os.makedirs("temp_test_files", exist_ok=True)
#     resume_path = "temp_test_files/resume.txt"
#     source1_path = "temp_test_files/source1.txt"
#     source2_path = "temp_test_files/unsupported.xyz" # Test unsupported
#     pdf_path = "temp_test_files/dummy.pdf"
#     docx_path = "temp_test_files/dummy.docx"

#     with open(resume_path, "w") as f: f.write("This is a resume with skills and experience. Contact: test@example.com")
#     with open(source1_path, "w") as f: f.write("This is a source document with project details.")
#     with open(source2_path, "w") as f: f.write("This is an unsupported file.")

#     from pypdf import PdfWriter # Ensure pypdf is installed
#     writer = PdfWriter()
#     # Add a page with some text (pypdf doesn't directly add text easily, this is a workaround)
#     # For a real test, use a PDF with actual text. Here, we create a blank one.
#     # For testing text extraction, a PDF file with actual text content is needed.
#     # This dummy PDF will likely result in empty text from pypdf.
#     # Consider adding a small, simple PDF with text to the repo for testing.
#     page = writer.add_blank_page(width=210, height=297)
#     with open(pdf_path, "wb") as f: writer.write(f)
#     print(f"Created dummy PDF: {pdf_path}") # Not prepending trace_id, example code

#     from docx import Document as DocxDoc # Ensure python-docx is installed
#     doc = DocxDoc()
#     doc.add_paragraph("This is a paragraph in a dummy docx file.")
#     doc.add_paragraph("Another paragraph with some more text.")
#     doc.save(docx_path)
#     print(f"Created dummy DOCX: {docx_path}") # Not prepending trace_id, example code


#     initial_state = GraphState(
#         client_id="test_client_123",
#         resume_file_path=resume_path,
#         source_document_paths=[source1_path, pdf_path, docx_path, source2_path, "temp_test_files/nonexistent.doc"],
#         # Initialize all keys expected by the graph state
#         source_document_texts=[],
#         resume_text=None,
#         roles_draft=None,
#         synced_roles_count=None, # Added
#         error_message=None,
#         current_node_error=None,
#         trace_id=None # Added
#     )

#     # To run the graph (uncomment when graph is fully defined and compiled)
#     # output = app.invoke(initial_state)
#     # print("\n--- Final State (after running graph) ---")
#     # for key, value in output.items():
#     #    if isinstance(value, list) and len(value) > 1 and isinstance(value[0], str) and len(value[0]) > 100:
#     #        print(f"{key}: [List of long texts, first 100 chars of first item: '{value[0][:100]}...']")
#     #    else:
#     #        print(f"{key}: {value}")

#     # To test a single node:
#     print("\n--- Testing extract_node directly ---")
#     extract_output = extract_node(initial_state)
#     print("\n--- Output of extract_node ---")
#     for key, value in extract_output.items():
#         if key == "source_document_texts" and isinstance(value, list):
#             print(f"{key}:")
#             for i, text in enumerate(value):
#                 print(f"  Doc {i}: '{text[:200]}...' (length: {len(text)})")
#         elif key == "resume_text" and isinstance(value, str):
#             print(f"{key}: '{value[:200]}...' (length: {len(value)})")
#         else:
#             print(f"{key}: {value}")

#     current_state = initial_state.copy()
#     current_state.update(extract_output) # Apply changes from extract_node

#     print("\n--- Testing parse_roles_node directly ---")
#     parse_output = parse_roles_node(current_state)
#     print("\n--- Output of parse_roles_node ---")
#     for key, value in parse_output.items(): print(f"{key}: {value}")

#     current_state.update(parse_output)

#     print("\n--- Testing sync_roles_node directly ---")
#     sync_output = sync_roles_node(current_state)
#     print("\n--- Output of sync_roles_node ---")
#     for key, value in sync_output.items(): print(f"{key}: {value}")

#     # Clean up dummy files
#     os.remove(resume_path)
#     os.remove(source1_path)
#     os.remove(source2_path)
#     os.remove(pdf_path)
#     os.remove(docx_path)
#     os.rmdir("temp_test_files")
